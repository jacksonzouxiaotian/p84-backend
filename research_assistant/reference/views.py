from flask import Blueprint, request, jsonify, send_file
from flask_jwt_extended import jwt_required, get_jwt_identity
from research_assistant.reference.models import Reference
from research_assistant.extensions import db

from io import BytesIO
from docx import Document
from docx.shared import Pt

# Optional: .bib parsing (install bibtexparser)
import bibtexparser
from bibtexparser.bparser import BibTexParser
from bibtexparser.customization import convert_to_unicode, homogenize_latex_encoding

bp = Blueprint("reference", __name__, url_prefix="/references")


def _normalize_authors_from_bib(entry):
    """Convert BibTeX 'author' to 'Lastname, F.; Foo, B.'"""
    raw = (entry.get("author") or "").strip()
    if not raw:
        return ""
    parts = [p.strip() for p in raw.split(" and ") if p.strip()]
    norm = []
    for p in parts:
        if "," in p:
            last, first = [x.strip() for x in p.split(",", 1)]
        else:
            tokens = p.split()
            last = tokens[-1]
            first = " ".join(tokens[:-1])
        initials = "".join([s[0].upper() + "." for s in first.split() if s])
        norm.append(f"{last}, {initials}".strip())
    return "; ".join(norm)


def _load_bib_entries(file_stream):
    parser = BibTexParser(common_strings=True)
    parser.customization = lambda rec: convert_to_unicode(homogenize_latex_encoding(rec))
    bib_db = bibtexparser.load(file_stream, parser=parser)
    return bib_db.entries


# -------------------- CRUD --------------------

@bp.route("/", methods=["POST"])
@jwt_required()
def add_reference():
    data = request.get_json() or {}
    title = data.get("title")
    authors = data.get("authors")
    year = data.get("year")
    source = data.get("source")
    user_id = int(get_jwt_identity())

    if not title or not authors or not year:
        return jsonify({"error": "Missing fields"}), 400

    ref = Reference(
        title=title,
        authors=authors,
        year=str(year),
        source=source,
        user_id=user_id,
    )
    db.session.add(ref)
    db.session.commit()
    return jsonify(ref.to_dict()), 201


@bp.route("/", methods=["GET"])
@jwt_required()
def list_references():
    user_id = int(get_jwt_identity())
    sort_by = request.args.get("sort_by", "created_at")
    allowed = {"created_at", "title", "year", "id"}
    if sort_by not in allowed:
        sort_by = "created_at"

    refs = (
        Reference.query
        .filter_by(user_id=user_id)
        .order_by(getattr(Reference, sort_by))
        .all()
    )
    return jsonify([ref.to_dict() for ref in refs])


@bp.route("/<int:ref_id>", methods=["PUT"])
@jwt_required()
def update_reference(ref_id):
    ref = Reference.query.get_or_404(ref_id)

    if ref.user_id != int(get_jwt_identity()):
        return jsonify({"error": "Unauthorized"}), 403

    data = request.get_json() or {}
    for field in ["title", "authors", "year", "source", "completed"]:
        if field in data:
            setattr(ref, field, data[field])

    db.session.commit()
    return jsonify(ref.to_dict())


@bp.route("/<int:ref_id>", methods=["DELETE"])
@jwt_required()
def delete_reference(ref_id):
    user_id = int(get_jwt_identity())
    ref = Reference.query.filter_by(id=ref_id, user_id=user_id).first_or_404()
    db.session.delete(ref)
    try:
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500
    return jsonify({"msg": "Deleted successfully"})



# -------------------- Upload .bib  --------------------

@bp.route("/upload_bib", methods=["POST"])
@jwt_required()
def upload_bib():
    """
    Strict to current Reference model:
      Only store: title, authors, year, source(='journal'), user_id
      Only handle @article entries.
    """
    user_id = int(get_jwt_identity())
    f = request.files.get("file")
    if not f:
        return jsonify({"error": "No file uploaded"}), 400

    try:
        entries = _load_bib_entries(f.stream)
    except Exception as e:
        return jsonify({"error": f"Failed to parse .bib: {e}"}), 400

    created = []
    for e in entries:
        if (e.get("ENTRYTYPE") or "").lower() != "article":
            continue

        title = (e.get("title") or "").strip("{} ")
        authors = _normalize_authors_from_bib(e)
        year = (e.get("year") or "").strip()
        if not (title and authors and year):
            continue

        ref = Reference(
            user_id=user_id,
            title=title,
            authors=authors,
            year=str(year),
            source="journal",
        )
        db.session.add(ref)
        db.session.flush()
        created.append(ref.to_dict())

    db.session.commit()
    return jsonify({"count": len(created), "created": created}), 201


# -------------------- Generate .docx citation  --------------------

@bp.route("/<int:ref_id>/cite", methods=["GET"])
@jwt_required()
def generate_citation_api(ref_id):
    """
    Generate a .docx file for the citation and trigger download.
    Supported styles: APA | CHICAGO | MLA
    """
    style = (request.args.get("style", "APA") or "APA").upper()
    if style not in {"APA", "CHICAGO", "MLA"}:
        return jsonify({"error": f"Unsupported style: {style}"}), 400

    ref = Reference.query.get_or_404(ref_id)
    if ref.user_id != int(get_jwt_identity()):
        return jsonify({"error": "Unauthorized"}), 403

    try:
        file_bytes, download_name = build_docx_citation(ref, style)
    except Exception as e:
        return jsonify({"error": f"Failed to build citation: {e}"}), 400

    return send_file(
        file_bytes,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=download_name,
        max_age=0,
    )



# -------------------- Docx builder & utils --------------------


def build_docx_citation(ref, style: str):
    """
    Build a single citation into a .docx for the given style.
    With current model, we only have: authors, year, title.
    APA  : Authors (Year). Title.
    Chicago: Authors. "Title." Year.
    MLA  : Authors. "Title." Year.
    """
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.size = Pt(11)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)

    authors_raw = getattr(ref, "authors", "") or ""
    year = getattr(ref, "year", "") or ""
    title = (getattr(ref, "title", "") or "").strip()

    if style == "APA":
        authors_text = format_authors_apa(authors_raw)
        # Authors (Year). Title.
        if authors_text:
            add_run(p, authors_text + " ")
        if year:
            add_run(p, f"({year}). ")
        if title:
            add_run(p, f"{title}. ")

    elif style == "CHICAGO":
        authors_text = format_authors_chicago(authors_raw)
        # Authors. "Title." Year.
        if authors_text:
            add_run(p, authors_text + ". ")
        if title:
            add_run(p, f"\"{title}.\" ")
        if year:
            add_run(p, f"{year}. ")

    elif style == "MLA":
        authors_text = format_authors_mla(authors_raw)
        # Authors. "Title." Year.
        if authors_text:
            add_run(p, authors_text + ". ")
        if title:
            add_run(p, f"\"{title}.\" ")
        if year:
            add_run(p, f"{year}. ")

    # 生成内存文件 & 下载名
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)

    first_author = extract_first_author(format_authors_apa(authors_raw))
    year_str = str(year or "")
    name = f"{first_author}_{year_str}_{style}.docx".strip("_").replace("__", "_")
    safe = "".join(c if c.isalnum() or c in "._-" else "_" for c in name)
    return bio, (safe or f"citation_{style}.docx")


def add_run(paragraph, text, italic=False):
    r = paragraph.add_run(text)
    if italic:
        r.italic = True
    return r


def format_authors_apa(authors):
    if not authors:
        return ""
    if isinstance(authors, str):
        items = [a.strip() for a in authors.split(";") if a.strip()]
    elif isinstance(authors, (list, tuple)):
        items = [str(a).strip() for a in authors if str(a).strip()]
    else:
        return str(authors)

    n = len(items)
    if n == 0:
        return ""
    if n == 1:
        return items[0]
    if n == 2:
        return f"{items[0]} & {items[1]}"
    return f"{', '.join(items[:-1])}, & {items[-1]}"

def _split_author_item(item: str):
    """
    Input like: 'Lastname, F.' or 'Lastname, First Middle'
    Return (last, first_part) without surrounding spaces/dots normalization.
    """
    item = (item or "").strip()
    if not item:
        return "", ""
    if "," in item:
        last, first = [x.strip() for x in item.split(",", 1)]
    else:
        # fallback: space-split -> last token as last name
        tokens = item.split()
        if not tokens:
            return "", ""
        last = tokens[-1]
        first = " ".join(tokens[:-1])
    return last, first


def _to_first_last(item: str) -> str:
    """'Lastname, F.' -> 'F. Lastname'  (for Chicago/MLA subsequent authors)"""
    last, first = _split_author_item(item)
    return (f"{first} {last}").strip()


def _to_last_first(item: str) -> str:
    """Normalize to 'Lastname, First/Initials' (for leading author in MLA/Chicago)."""
    last, first = _split_author_item(item)
    if not last and not first:
        return item.strip()
    return f"{last}, {first}".strip().strip(",")


def _authors_list(authors: str):
    """Return list from 'A, B.; C, D.' -> ['A, B.', 'C, D.']"""
    if not authors:
        return []
    return [a.strip() for a in authors.split(";") if a.strip()]


def format_authors_chicago(authors: str) -> str:
    """
    Chicago bibliography: first author as 'Last, First', subsequent as 'First Last'.
    Join with comma, final 'and'.
    """
    items = _authors_list(authors)
    n = len(items)
    if n == 0:
        return ""
    if n == 1:
        return _to_last_first(items[0])

    first = _to_last_first(items[0])
    rest = [_to_first_last(x) for x in items[1:]]
    if len(rest) == 1:
        return f"{first} and {rest[0]}"
    return f"{first}, {', '.join(rest[:-1])}, and {rest[-1]}"


def format_authors_mla(authors: str) -> str:
    """
    MLA 9:
      1 author: 'Last, First'
      2 authors: 'Last, First, and First Last'
      3+ authors: 'Last, First, et al.'
    """
    items = _authors_list(authors)
    n = len(items)
    if n == 0:
        return ""
    if n == 1:
        return _to_last_first(items[0])
    if n == 2:
        first = _to_last_first(items[0])
        second = _to_first_last(items[1])
        return f"{first}, and {second}"
    # 3+
    first = _to_last_first(items[0])
    return f"{first}, et al."


def strip_doi_prefix(doi: str) -> str:
    doi = (doi or "").strip()
    for pref in ("https://doi.org/", "http://doi.org/", "doi:", "DOI:"):
        if doi.lower().startswith(pref):
            return doi[len(pref):]
    return doi


def extract_first_author(authors_text: str) -> str:
    if not authors_text:
        return "citation"
    return authors_text.split("&")[0].split(",")[0].strip() or "citation"
